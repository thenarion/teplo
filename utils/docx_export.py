"""
Экспорт результатов теплового баланса в DOCX.
Использует python-docx.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io


def _make_charts(res) -> dict:
    """
    Генерирует plotly-графики и возвращает PNG-байты.
    Возвращает dict с ключами: mass, heat, stages (bytes или None).
    """
    try:
        import plotly.graph_objects as go
        import plotly.express as px
        import pandas as pd
    except ImportError:
        return {"mass": None, "heat": None, "stages": None}

    charts = {"mass": None, "heat": None, "stages": None}
    inp = res.input
    b = res.burnout

    # 1. Круговая диаграмма состава топлива
    try:
        mass_df = pd.DataFrame({
            "Компонент": ["Вода", "Зола", "Горючая масса"],
            "Масса, кг/ч": [res.water_mass, res.ash_mass, res.combustible_mass],
        })
        fig = px.pie(mass_df, values="Масса, кг/ч", names="Компонент",
                     title="Состав топлива", hole=0.3,
                     color="Компонент",
                     color_discrete_map={"Вода": "#3498db", "Зола": "#95a5a6", "Горючая масса": "#e74c3c"})
        fig.update_layout(width=600, height=400, font=dict(size=12),
                         template="plotly_white")
        charts["mass"] = fig.to_image(format="png", engine="kaleido")
    except Exception:
        pass

    # 2. Столбчатая диаграмма теплового баланса
    try:
        fig = go.Figure()
        fig.add_trace(go.Bar(x=["Приход"], y=[res.q_fuel_actual],
                             name="Тепло от топлива", marker_color="green"))
        fig.add_trace(go.Bar(x=["Приход"], y=[inp.burner_power],
                             name="Тепло от горелки", marker_color="lightgreen"))
        fig.add_trace(go.Bar(x=["Расход"], y=[res.q_flue_gas],
                             name="Уходящие газы", marker_color="red"))
        fig.add_trace(go.Bar(x=["Расход"], y=[res.q_wall],
                             name="Потери через футеровку", marker_color="orange"))
        fig.add_trace(go.Bar(x=["Расход"], y=[res.q_ash],
                             name="Потери с золой", marker_color="yellow"))
        fig.add_trace(go.Bar(x=["Расход"], y=[res.q_useful_with_burner],
                             name="Полезное тепло", marker_color="blue"))
        fig.update_layout(title="Тепловой баланс, МВт", barmode="stack",
                          yaxis_title="МВт", width=700, height=400, font=dict(size=12))
        charts["heat"] = fig.to_image(format="png", engine="kaleido")
    except Exception:
        pass

    # 3. Диаграмма стадий выгорания
    if b:
        try:
            stages_df = pd.DataFrame({
                "Стадия": ["Сушка", "Нагрев", "Горение", "Дожигание"],
                "Время, мин": [b.t_drying, b.t_heating, b.t_combustion, b.t_burnout],
            })
            fig = px.bar(stages_df, x="Стадия", y="Время, мин",
                         title="Время по стадиям выгорания", color="Стадия",
                         color_discrete_sequence=["#FF6B6B", "#FFA500", "#4ECDC4", "#45B7D1"])
            fig.add_hline(y=b.residence_time, line_dash="dash", line_color="red",
                          annotation_text=f"Факт. время: {b.residence_time:.0f} мин")
            fig.add_hline(y=b.t_required, line_dash="dot", line_color="orange",
                          annotation_text=f"Необходимое: {b.t_required:.0f} мин")
            fig.update_layout(width=700, height=400, font=dict(size=12))
            charts["stages"] = fig.to_image(format="png", engine="kaleido")
        except Exception:
            pass

    return charts


def _add_page_number(paragraph):
    """Добавляет поле номера страницы в параграф."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def _add_table(doc, headers, rows, col_widths_cm=None):
    """Добавляет таблицу с заголовками и данными."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Заголовки
    hdr_cells = table.rows[0].cells
    for i, name in enumerate(headers):
        hdr_cells[i].text = name
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)

    # Данные
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, value in enumerate(row):
            cells[c_idx].text = str(value)
            for p in cells[c_idx].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    # Ширина столбцов
    if col_widths_cm:
        for i, width in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # отступ после таблицы


def _add_image(doc, img_bytes, width_cm=15):
    """Добавляет изображение из байтов."""
    stream = io.BytesIO(img_bytes)
    doc.add_picture(stream, width=Cm(width_cm))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def export_to_docx(res, summary_table: list, flue_gas_params: list) -> bytes:
    """
    Экспортирует результаты теплового баланса в DOCX.
    Структура аналогична PDF-экспорту.
    """
    doc = Document()

    # Настройка страницы A4
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Стили
    style_normal = doc.styles["Normal"]
    style_normal.font.size = Pt(11)
    style_normal.font.name = "Calibri"

    inp = res.input
    b = res.burnout

    # Генерируем графики
    charts = _make_charts(res)

    # Заголовок документа
    title = doc.add_paragraph("Тепловой баланс сжигания топлива", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(
        f"Дата расчёта: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
        style="Subtitle",
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # отступ

    # =========================================================
    # 1. ИСХОДНЫЕ ДАННЫЕ
    # =========================================================
    doc.add_heading("1. Исходные данные", level=1)

    input_data = [
        ["Подача влажного топлива", f"{inp.fuel_feed:.0f} кг/ч"],
        ["Влажность", f"{inp.moisture*100:.0f}%"],
        ["Низшая теплота сгорания (Q_net)", f"{inp.q_net_ar:.2f} МДж/кг"],
        ["База теплоты сгорания", {
            "as_received": "Рабочая масса (с водой и золой)",
            "dry": "Сухая масса (без воды)",
            "daf": "Горючая масса (без воды и золы)",
        }.get(inp.q_basis, inp.q_basis)],
        ["Зольность на рабочую массу", f"{inp.ash_content*100:.0f}%"],
        ["Коэффициент избытка воздуха (α)", f"{inp.excess_air:.2f}"],
        ["Температура дымовых газов на выходе", f"{inp.flue_gas_temp:.0f}°C"],
        ["Температура наружного воздуха", f"{inp.ambient_temp:.0f}°C"],
        ["Мощность горелки (макс)", f"{inp.burner_power:.1f} МВт"],
        ["Длина барабана", f"{inp.drum_length:.1f} м"],
        ["Диаметр барабана", f"{inp.drum_diameter:.1f} м"],
        ["Время пребывания отхода", f"{b.residence_time:.0f} мин"],
        ["Насыпная плотность отхода", f"{inp.bulk_density:.0f} кг/м³"],
    ]

    _add_table(doc, ["Параметр", "Значение"], input_data, col_widths_cm=[10, 7])

    # =========================================================
    # 2. МАССОВЫЙ БАЛАНС
    # =========================================================
    doc.add_heading("2. Массовый баланс", level=1)

    mass_data = [
        ["Влажный помёт (всего)", f"{res.fuel_feed:.0f} кг/ч", "100%"],
        ["Вода", f"{res.water_mass:.0f} кг/ч", f"{res.water_mass/res.fuel_feed*100:.0f}%"],
        ["Сухое вещество", f"{res.dry_mass:.0f} кг/ч", f"{res.dry_mass/res.fuel_feed*100:.0f}%"],
        ["Зола", f"{res.ash_mass:.0f} кг/ч", f"{res.ash_mass/res.fuel_feed*100:.0f}%"],
        ["Горючая масса", f"{res.combustible_mass:.0f} кг/ч", f"{res.combustible_mass/res.fuel_feed*100:.0f}%"],
    ]

    _add_table(doc, ["Компонент", "Масса", "Доля"], mass_data, col_widths_cm=[6, 5, 4])

    if charts["mass"]:
        _add_image(doc, charts["mass"], width_cm=12)

    # =========================================================
    # 3. ТЕПЛОТА СГОРАНИЯ
    # =========================================================
    doc.add_heading("3. Теплота сгорания", level=1)

    heat_data = [
        ["Низшая теплота сгорания (Q_net)", f"{res.q_net_ar:.2f} МДж/кг"],
        ["Тепловыделение от топлива", f"{res.q_fuel_actual:.3f} МВт"],
    ]

    _add_table(doc, ["Параметр", "Значение"], heat_data, col_widths_cm=[10, 7])

    # =========================================================
    # 4. РАСХОД ВОЗДУХА И ДЫМОВЫХ ГАЗОВ
    # =========================================================
    doc.add_heading("4. Расход воздуха и дымовых газов", level=1)

    air_data = [
        ["Теоретический объём воздуха (на кг)", f"{res.v_air_theoretical_per_kg:.2f} Нм³/кг"],
        [f"Фактический объём воздуха (α={inp.excess_air:.2f})", f"{res.v_air_actual:.0f} Нм³/ч"],
        ["Масса воздуха", f"{res.m_air:.0f} кг/ч"],
        ["Объём дымовых газов (н.у.)", f"{res.v_flue:.0f} Нм³/ч"],
        ["Масса дымовых газов", f"{res.m_flue:.0f} кг/ч"],
    ]

    _add_table(doc, ["Параметр", "Значение"], air_data, col_widths_cm=[10, 7])

    # =========================================================
    # 5. СВОДНАЯ ТАБЛИЦА ТЕПЛОВОГО БАЛАНСА
    # =========================================================
    doc.add_heading("5. Сводная таблица теплового баланса", level=1)

    balance_data = [[row["Статья"], row["МВт"], row["%"]] for row in summary_table]

    _add_table(doc, ["Статья", "МВт", "%"], balance_data, col_widths_cm=[9, 4, 4])

    if charts["heat"]:
        _add_image(doc, charts["heat"], width_cm=15)

    # =========================================================
    # 6. ПОЛНОТА ВЫГОРАНИЯ
    # =========================================================
    doc.add_heading("6. Полнота выгорания", level=1)

    burnout_data = [
        ["Объём барабана", f"{b.drum_volume:.2f} м³"],
        ["Масса отхода в барабане", f"{b.mass_in_drum:.0f} кг"],
        ["Степень заполнения", f"{b.fill_ratio*100:.1f}%"],
        ["Время сушки", f"{b.t_drying:.1f} мин"],
        ["Время нагрева", f"{b.t_heating:.1f} мин"],
        ["Время горения", f"{b.t_combustion:.1f} мин"],
        ["Время дожигания", f"{b.t_burnout:.1f} мин"],
        ["Необходимое время выгорания", f"{b.t_required:.1f} мин"],
        ["Фактическое время пребывания", f"{b.residence_time:.1f} мин"],
        ["Коэффициент запаса времени", f"{b.time_ratio:.2f}"],
        ["Полнота выгорания", f"{b.burnout_efficiency*100:.1f}%"],
        ["Удельная тепловая нагрузка", f"{b.heat_load:.0f} кВт/м³"],
    ]

    _add_table(doc, ["Параметр", "Значение"], burnout_data, col_widths_cm=[10, 7])

    if charts["stages"]:
        _add_image(doc, charts["stages"], width_cm=15)

    # =========================================================
    # 7. ВЫВОДЫ
    # =========================================================
    doc.add_heading("7. Выводы", level=1)

    conclusions = [
        f"Автотермичность: {'Да' if res.q_useful_no_burner > 0 else 'Нет'}",
        f"Тепловыделение от топлива: {res.q_fuel_actual:.3f} МВт",
        f"Полезное тепло без горелки: {res.q_useful_no_burner:.3f} МВт",
        f"КПД установки: {res.efficiency_with_burner*100:.1f}%",
        "",
        f"Время пребывания: {b.residence_time:.1f} мин (необходимо {b.t_required:.1f} мин)",
        f"Запас времени: ×{b.time_ratio:.2f}",
        f"Время пребывания достаточное: {'Да' if b.time_ok else 'Нет'}",
        "",
        f"Степень заполнения: {b.fill_ratio*100:.1f}% (допустимая {inp.max_fill_ratio*100:.0f}%)",
        f"Степень заполнения в норме: {'Да' if b.fill_ratio_ok else 'Нет'}",
        "",
        f"Тепловая нагрузка: {b.heat_load:.0f} кВт/м³ (допустимая {inp.max_heat_load:.0f} кВт/м³)",
        f"Тепловая нагрузка в норме: {'Да' if b.heat_load_ok else 'Нет'}",
        "",
        f"Полнота выгорания: {b.burnout_efficiency*100:.1f}%",
        f"Оценка: {b.burnout_status_ru}",
    ]

    for line in conclusions:
        doc.add_paragraph(line)

    doc.add_paragraph()  # отступ

    # Общий вывод
    if b.overall_ok:
        p = doc.add_paragraph()
        run = p.add_run("ОБЩИЙ ВЫВОД: ")
        run.bold = True
        p.add_run("Отход успевает полностью выгореть. Установка работает в штатном режиме.")
    else:
        reasons = []
        if not b.time_ok:
            reasons.append("недостаточное время пребывания")
        if not b.fill_ratio_ok:
            reasons.append("превышена степень заполнения")
        if not b.heat_load_ok:
            reasons.append("превышена тепловая нагрузка")
        if b.burnout_efficiency < 0.90:
            reasons.append(f"полнота выгорания ниже 90% (фактически {b.burnout_efficiency*100:.1f}%)")

        p = doc.add_paragraph()
        run = p.add_run("ОБЩИЙ ВЫВОД: ")
        run.bold = True
        p.add_run("Отход НЕ успевает полностью выгореть.")

        p = doc.add_paragraph()
        run = p.add_run("Причины: ")
        run.bold = True
        p.add_run(f"{', '.join(reasons)}.")

    # Нижний колонтитул с номером страницы
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.style = doc.styles["Normal"]
    run = footer.add_run("Страница ")
    run.font.size = Pt(8)
    _add_page_number(footer)

    # Сохраняем в байты
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
